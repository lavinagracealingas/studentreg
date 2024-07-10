import streamlit as st
import calendar
from datetime import datetime
import pandas as pd
import numpy as np
from streamlit_option_menu import option_menu
import pickle
from pathlib import Path
import streamlit_authenticator as stauth
from streamlit_pandas_profiling import st_profile_report
import sqlite3
import plotly.express as px
import docx 
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64


conn = sqlite3.connect('studentmonitor.db', check_same_thread=False)
cur = conn.cursor()

def calculate_rates(academic_year):
    cur.execute("SELECT COUNT(StudentID) FROM academicrecords WHERE AcademicYear <= ?", (academic_year,))
    student_total = cur.fetchone()[0]

    cur.execute("SELECT COUNT(StudentID) FROM academicrecords WHERE YearLevel = '1' AND AcademicYear = ?", (academic_year,))
    initial_cohort_size = cur.fetchone()[0]

    cur.execute("SELECT COUNT(StudentID) FROM academicrecords WHERE YearLevel IN ('1', '2', '3', '4') AND AcademicYear = ?", (academic_year,))
    current_students = cur.fetchone()[0]
    retention_rate = (current_students / initial_cohort_size) * 100 if initial_cohort_size > 0 else 0

    cur.execute("SELECT COUNT(StudentID) FROM academicrecords WHERE ScholasticStatus = 'Graduate' AND AcademicYear = ?", (academic_year,))
    graduate_count = cur.fetchone()[0]
    completion_rate = (graduate_count / student_total) * 100 if student_total > 0 else 0

    cur.execute("SELECT COUNT(StudentID) FROM promotion WHERE PromotionStatus = '1' AND AcademicYear = ?", (academic_year,))
    promotion_count = cur.fetchone()[0]
    promotion_rate = (promotion_count / student_total) * 100 if student_total > 0 else 0

    cur.execute("SELECT COUNT(StudentID) FROM courseassignment WHERE GradeStatus = 'Failed' AND AcademicYear = ?", (academic_year,))
    fail_count = cur.fetchone()[0]
    failure_rate = (fail_count / student_total) * 100 if student_total > 0 else 0

    cur.execute("SELECT COUNT(StudentID) FROM academicrecords WHERE ScholasticStatus = 'Dropped' AND AcademicYear = ?", (academic_year,))
    dropout_count = cur.fetchone()[0]
    dropout_rate = (dropout_count / student_total) * 100 if student_total > 0 else 0

    return {
        "student_total": student_total,
        "retention_rate": retention_rate,
        "completion_rate": completion_rate,
        "promotion_rate": promotion_rate,
        "failure_rate": failure_rate,
        "dropout_rate": dropout_rate,
        "fail_count": fail_count,
        "dropout_count": dropout_count
    }

def calculate_gpa(student_id, year_level, semester):
    cur.execute("""
        SELECT ca.Grade, ca.FinalGrade, p.Units
        FROM courseassignment ca
        JOIN prospectus p ON ca.CourseCode = p.CourseCode
        WHERE ca.StudentID = ? AND ca.YearLevel = ? AND ca.Semester = ? 
          AND ca.Grade NOT IN ('W', 'P', 'F', 'INPROG') 
          AND ca.CourseCode NOT IN ('NST001', 'NST002')
    """, (student_id, year_level, semester))
    grades = cur.fetchall()

    if not grades:
        return None

    total_units = 0
    weighted_sum = 0.0

    for grade in grades:
        initial_grade = grade[0]
        final_grade = grade[1]
        units = grade[2]

        if initial_grade in ['INC', 'INPROG']:
            if final_grade and final_grade.strip() != '':
                try:
                    weighted_sum += float(final_grade) * units
                    total_units += units
                except ValueError:
                    continue  # Skip this grade if final grade cannot be converted to float
        else:
            try:
                weighted_sum += float(initial_grade) * units
                total_units += units
            except ValueError:
                continue  # Skip this grade if initial grade cannot be converted to float

    if total_units > 0:
        gpa_value = round(weighted_sum / total_units, 2)
    else:
        gpa_value = 0.0

    return gpa_value

def calculate_awardees(year_level, semester):
    cur.execute("""
        SELECT DISTINCT StudentID FROM courseassignment
        WHERE YearLevel = ? AND Semester = ?
    """, (year_level, semester))
    student_ids = [row[0] for row in cur.fetchall()]

    rl_count = 0
    cl_count = 0
    dl_count = 0

    for student_id in student_ids:
        gpa = calculate_gpa(student_id, year_level, semester)
        if gpa is not None:
            if 1.0 <= gpa <= 1.20:
                rl_count += 1
            elif 1.21 <= gpa <= 1.45:
                cl_count += 1
            elif 1.46 <= gpa <= 1.75:
                dl_count += 1

    return {
        "rl_count": rl_count,
        "cl_count": cl_count,
        "dl_count": dl_count
    }

def calculate_counts(year_level, semester):
    cur.execute("SELECT COUNT(StudentID) FROM courseassignment WHERE Grade = 'INC' AND YearLevel = ? AND Semester = ?", (year_level, semester))
    inc_count = cur.fetchone()[0]

    cur.execute("SELECT COUNT(StudentID) FROM academicrecords WHERE ScholasticStatus = 'Withdrawn' AND YearLevel = ? AND Semester = ?", (year_level, semester))
    withdrawn_count = cur.fetchone()[0]

    cur.execute("SELECT COUNT(StudentID) FROM courseassignment WHERE GradeStatus = 'Failed' AND YearLevel = ? AND Semester = ?", (year_level, semester))
    fail_count = cur.fetchone()[0]

    return {
        "inc_count": inc_count,
        "withdrawn_count": withdrawn_count,
        "fail_count": fail_count,
        **calculate_awardees(year_level, semester)
    }

def calculate_cgpa(student_id, year_level, semester):
    cur.execute("""
        SELECT ca.StudentID, ca.CourseCode, p.Units, ca.Grade, ca.FinalGrade
        FROM courseassignment ca
        JOIN prospectus p ON ca.CourseCode = p.CourseCode
        WHERE ca.StudentID = ? AND ca.YearLevel = ? AND ca.Semester = ?
    """, (student_id, year_level, semester))
    
    rows = cur.fetchall()
    total_units = 0
    total_grade_points = 0
    
    for row in rows:
        units = float(row[2])  # Convert units to float
        grade = row[3]  # Grade from courseassignment table
        final_grade = row[4]  # FinalGrade from courseassignment table
        
        # Consider only final grades for CGPA calculation
        if final_grade is not None and final_grade.strip():  # Check if final_grade is not empty or None
            try:
                grade_points = float(final_grade) * units
                total_grade_points += grade_points
                total_units += units
            except ValueError:
                continue  # Skip this grade if final grade cannot be converted to float
    
    if total_units == 0:
        return None  # No units found for the student
    
    cgpa = total_grade_points / total_units
    return cgpa


def get_initial_grade_value(initial_grade, final_grade):
    if initial_grade in ["1.00", "1.25", "1.50", "1.75", "2.00", "2.25", "2.50", "2.75", "3.00"]:
        return float(initial_grade)
    elif initial_grade in ["INC", "INPROG"] and final_grade in ["1.00", "1.25", "1.50", "1.75", "2.00", "2.25", "2.50", "2.75", "3.00"]:
        return float(final_grade)
    elif initial_grade == "5.00":
        return 5.00
    else:
        return None
def calc_gpa(grades_df):
    gpa_data = []
    for student_id, group in grades_df.groupby('StudentID'):
        valid_grades = group[~group['CourseCode'].isin(['NST001', 'NST002'])]
        valid_grades['GradePoint'] = valid_grades.apply(lambda row: get_initial_grade_value(row['Grade'], row['FinalGrade']), axis=1)
        valid_grades.dropna(subset=['GradePoint'], inplace=True)
        total_units = valid_grades['Units'].sum()
        weighted_sum = (valid_grades['Units'] * valid_grades['GradePoint']).sum()
        if total_units > 0:
            gpa = weighted_sum / total_units
            gpa_data.append((student_id, gpa))
    return pd.DataFrame(gpa_data, columns=['StudentID', 'GPA'])

def calc_cgpa(grades_df):
    cgpa_data = []
    for student_id, group in grades_df.groupby('StudentID'):
        running_total_units = 0
        running_weighted_sum = 0
        valid_grades = group[~group['CourseCode'].isin(['NST001', 'NST002'])]
        valid_grades['GradePoint'] = valid_grades.apply(lambda row: get_initial_grade_value(row['Grade'], row['FinalGrade']), axis=1)
        valid_grades.dropna(subset=['GradePoint'], inplace=True)
        running_total_units += valid_grades['Units'].sum()
        running_weighted_sum += (valid_grades['Units'] * valid_grades['GradePoint']).sum()
        if running_total_units > 0:
            cgpa = round(running_weighted_sum / running_total_units, 5)
            cgpa_data.append((student_id, cgpa))
    return pd.DataFrame(cgpa_data, columns=['StudentID', 'CGPA'])

def get_all_student_grades():
    query = """
        SELECT ca.StudentID, ca.CourseCode, p.Units, ca.Grade, ca.FinalGrade, ca.YearLevel, ca.Semester
        FROM courseassignment ca
        JOIN prospectus p ON ca.CourseCode = p.CourseCode
    """
    grade_df = pd.read_sql_query(query, conn)
    return grade_df

def calculate_average_gpa_cgpa_all():
    final_grades_df = get_all_student_grades()
    all_year_levels = final_grades_df['YearLevel'].unique()
    all_semesters = final_grades_df['Semester'].unique()
    
    avg_gpa_cgpa_data = []
    
    for year_level in all_year_levels:
        for semester in all_semesters:
            selected_grades_df = final_grades_df[(final_grades_df['YearLevel'] == year_level) & 
                                                (final_grades_df['Semester'] == semester)]
            gpa_df = calc_gpa(selected_grades_df)
            cgpa_df = calc_cgpa(selected_grades_df)
            avg_gpa = gpa_df['GPA'].mean() if not gpa_df.empty else None
            avg_cgpa = cgpa_df['CGPA'].mean() if not cgpa_df.empty else None
            avg_gpa_cgpa_data.append((year_level, semester, avg_gpa, avg_cgpa))
    
    avg_gpa_cgpa_df = pd.DataFrame(avg_gpa_cgpa_data, columns=['YearLevel', 'Semester', 'AverageGPA', 'AverageCGPA'])
    return avg_gpa_cgpa_df

def get_course_data_with_status_counts(conn, year_level, semester):
            query = f"""
            SELECT p.CourseCode, p.CourseDesc, p.Units, ca.Semester, ca.YearLevel, 
            SUM(CASE WHEN ca.GradeStatus = 'Passed' THEN 1 ELSE 0 END) as PassedCount,
            SUM(CASE WHEN ca.GradeStatus = 'Failed' THEN 1 ELSE 0 END) as FailedCount,
            SUM(CASE WHEN ca.GradeStatus = 'Dropout' THEN 1 ELSE 0 END) as DroppedCount,
            SUM(CASE WHEN ca.GradeStatus = 'Withdrawn' THEN 1 ELSE 0 END) as WithdrawnCount,
            SUM(CASE WHEN ca.GradeStatus != 'Passed' THEN 1 ELSE 0 END) as RetakeCount
            FROM prospectus p
            LEFT JOIN courseassignment ca ON p.CourseCode = ca.CourseCode
            WHERE ca.YearLevel LIKE ? AND ca.Semester LIKE ?
            GROUP BY p.CourseCode, p.CourseDesc, p.Units, ca.Semester, ca.YearLevel
            """
            df = pd.read_sql_query(query, conn, params=(f'%{year_level}%', f'%{semester}%'))
            return df


def get_pdf_download_link(file_path):
    # Function to generate download link for PDF file
    with open(file_path, "rb") as f:
        pdf = f.read()
        b64_pdf = base64.b64encode(pdf).decode()
        href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="User Guide.pdf"><button>Download User Guide (PDF)</button></a>'
    return href

def app():
    st.subheader("Home", divider='red')
    tab1, tab2, tab3, tab4 = st.tabs(["About", "Counts", "Trends", "Adviser's Report"])

    with tab1:
        st.header("About Us")

        st.markdown("<br>", unsafe_allow_html=True)

        st.write("""
            <div style="text-align: justify;">
            We are a team of incoming 4th-year BS Statistics students currently 
            undertaking our summer internship at the <strong>MSU-IIT Premier Research Institute 
            of Science and Mathematics (PRISM)</strong>. As part of our training, we have been 
            tasked with developing a comprehensive student monitoring application. Our goal 
            is to create a tool that significantly aids academic advisers in managing and 
            tracking student performance, making their work more efficient and effective.
            </div>""", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        st.subheader("The App")
        st.write("""
            <div style="text-align: justify;">
            Our application is a tool designed to assist academic advisers in 
            managing and tracking student performance. It covers a broad scope of student 
            demographics, course enrollments, and grade evaluations. The primary objective 
            is to simplify the monitoring process, enabling advisers to provide timely 
            interventions and support to students when needed.
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        st.subheader("Features")
        st.markdown("""
            <div style="text-align: justify;">
            Our application offers comprehensive tracking, allowing advisers to 
            monitor student demographics, course enrollments, and grades efficiently. Data 
            management is made easy, enabling advisers to input, update, and manage student 
            information seamlessly. The application facilitates timely interventions by 
            quickly identifying and addressing academic concerns. Additionally, the 
            user-friendly interface ensures effortless navigation through the application, 
            with clearly labeled buttons and an organized layout designed to enhance the 
            user experience.
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        st.subheader("Contact Us")
        st.markdown("For any inquiries, support, or feedback, please get in touch with us:")
        st.write("- **Email:** waterlemonaide@gmail.com")
        st.write("- **Phone:** 09901234567")
        st.write("- **Address:** MSU-IIT Premier Research Institute of Science and Mathematics (PRISM), Iligan City, Lanao Del Norte")
        st.markdown("""
                    <div style="text-align: justify;">
                    We are here to assist you and ensure you have the best experience using our application. Your feedback is valuable to us and helps us improve continuously.
                    </div>
                    """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        st.subheader("User Guide")
        pdf_path = "User Guide.pdf" 
        download_link = get_pdf_download_link(pdf_path)
        st.markdown(download_link, unsafe_allow_html=True)

    with tab2:
        year_levels = ["1", "2", "3", "4"]
        semesters = ["1st Sem", "2nd Sem", "Summer"]
        col1, col2, col3, col4 = st.columns(4)
        selected_year_level = col1.selectbox("Select Year Level:", year_levels)
        selected_semester = col2.selectbox("Select Semester:", semesters)

        if selected_year_level and selected_semester:
            counts = calculate_counts(selected_year_level, selected_semester)

            st.subheader(f"Counts for {selected_year_level} Year: {selected_semester}")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("INC Grades", counts["inc_count"])
            with col2:
                st.metric("Withdrawn Students", counts["withdrawn_count"])
            with col3:
                st.metric("Failing Grades", counts["fail_count"])

            if any([counts["inc_count"], counts["withdrawn_count"], counts["fail_count"]]):
                fig = px.pie(names=["INC Grades", "Withdrawn Students", "Failing Grades"],
                            values=[counts["inc_count"], counts["withdrawn_count"], counts["fail_count"]],
                            title=f"Student Distribution for {selected_year_level} Year: {selected_semester}",
                            hole=0.5)
                st.plotly_chart(fig)

            st.divider()

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("RL Awardees", counts["rl_count"])
            with col2:
                st.metric("CL Awardees", counts["cl_count"])
            with col3:
                st.metric("DL Awardees", counts["dl_count"])

            if any([counts["rl_count"], counts["cl_count"], counts["dl_count"]]):
                fig = px.pie(names=["RL Awardees", "CL Awardees", "DL Awardees"],
                            values=[counts["rl_count"], counts["cl_count"], counts["dl_count"]],
                            title=f"Student Distribution for {selected_year_level} Year: {selected_semester}",
                            hole=0.5)
                st.plotly_chart(fig)

            # Calculate GPA distribution
            cur.execute("SELECT DISTINCT StudentID FROM academicrecords")
            student_ids = [row[0] for row in cur.fetchall()]
            below_25_gpa = 0
            above_25_gpa = 0
            below_25_cgpa = 0
            above_25_cgpa = 0

            for student_id in student_ids:
                gpa = calculate_gpa(student_id, selected_year_level, selected_semester)
                cgpa = calculate_cgpa(student_id, selected_year_level, selected_semester)
                
                if gpa is not None:
                    if gpa > 2.5:
                        below_25_gpa += 1
                    else:
                        above_25_gpa += 1
                
                if cgpa is not None:
                    if cgpa > 2.5:
                        below_25_cgpa += 1
                    else:
                        above_25_cgpa += 1

            # Display GPA and CGPA Distributions using Plotly
            st.subheader("GPA Distribution")
            gpa_fig = px.pie(names=["GPA Below 2.50", "GPA Above 2.50"],
                            values=[below_25_gpa, above_25_gpa],
                            title="GPA Distribution",
                            hole=0.5)
            st.plotly_chart(gpa_fig)

            st.subheader("CGPA Distribution")
            cgpa_fig = px.pie(names=["CGPA Below 2.50", "CGPA Above 2.50"],
                            values=[below_25_cgpa, above_25_cgpa],
                            title="CGPA Distribution",
                            hole=0.5)
            st.plotly_chart(cgpa_fig)

            st.divider()

            course_data_df = get_course_data_with_status_counts(conn, selected_year_level, selected_semester)
                
            if not course_data_df.empty:
                st.subheader(f'Course Data for {selected_year_level} Year Level, {selected_semester} Semester')
                # Drop the columns StudentID, YearLevel, and Semester before displaying
                course_display_df = course_data_df.drop(columns=['Units','YearLevel', 'Semester'])
                st.dataframe(course_display_df)
                
                # Visualization using Plotly
                fig = px.bar(course_data_df, x='CourseCode', y=['PassedCount', 'FailedCount', 'DroppedCount', 'WithdrawnCount', 'RetakeCount'],
                            title=f'Course Status Counts for {selected_year_level} Year Level, {selected_semester}')
                st.plotly_chart(fig)
            else:
                st.warning("No data found for the selected year level and semester.")
        

    with tab3:  
        # Select academic year
        cur.execute("SELECT DISTINCT AcademicYear FROM academicrecords ORDER BY AcademicYear")
        academic_years = [row[0] for row in cur.fetchall()]
        
        col1, col2, col3 = st.columns(3)
        selected_academic_year = col1.selectbox("Select Academic Year:", academic_years)

        if selected_academic_year:
            rates = calculate_rates(selected_academic_year)

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Number of students", rates["student_total"], delta=0, delta_color="normal")
                st.metric("Retention Rate", rates['retention_rate'], "%", delta_color="normal")
            with col2:
                st.metric("Completion Rate", rates['completion_rate'], "%", delta_color="normal")
                st.metric("Promotion Rate", rates['promotion_rate'], "%", delta_color="normal")
            with col3:
                st.metric("Failure Rate", rates['failure_rate'], "%", delta_color="normal")
                st.metric("Dropout Rate", rates['dropout_rate'], "%", delta_color="normal")

       # Fetch and calculate average GPA and CGPA
        avg_gpa_cgpa_df = calculate_average_gpa_cgpa_all()

        # Drop rows with None values to avoid plotting issues
        avg_gpa_cgpa_df.dropna(inplace=True)

        # Create the line graph for Average GPA
        fig_gpa = px.line(avg_gpa_cgpa_df, x='Semester', y='AverageGPA', 
                        color='YearLevel',
                        labels={'AverageGPA': 'Average GPA', 'Semester': 'Semester'},
                        title='Progression of Average GPA',
                        markers=True)

        fig_gpa.update_layout(
            xaxis_title="Semester",
            yaxis_title="Average GPA",
            legend_title_text='Year Level'
        )

        # Create the line graph for Average CGPA
        fig_cgpa = px.line(avg_gpa_cgpa_df, x='Semester', y='AverageCGPA', 
                        color='YearLevel',
                        labels={'AverageCGPA': 'Average CGPA', 'Semester': 'Semester'},
                        title='Progression of Average CGPA',
                        markers=True)

        fig_cgpa.update_layout(
            xaxis_title="Semester",
            yaxis_title="Average CGPA",
            legend_title_text='Year Level'
        )

        # Display the charts in the Streamlit app
        st.plotly_chart(fig_gpa)
        st.plotly_chart(fig_cgpa)

    with tab4:
        # Fetch and calculate average GPA and CGPA
        avg_gpa_cgpa_df = calculate_average_gpa_cgpa_all()

        # Check if DataFrame is empty before accessing its rows
        if not avg_gpa_cgpa_df.empty:
            average_gpa = avg_gpa_cgpa_df['AverageGPA'].values[0]
            average_cgpa = avg_gpa_cgpa_df['AverageCGPA'].values[0]
        else:
            st.error("No data available for average GPA and CGPA.")
            average_gpa = 0 
            average_cgpa = 0 

        cur.execute("SELECT DISTINCT AcademicYear FROM academicrecords ORDER BY AcademicYear")
        academic_year = [row[0] for row in cur.fetchall()]
        year_levels = ["1", "2", "3", "4"]
        semesters = ["1st Sem", "2nd Sem", "Summer"]

        col1, col2, col3 =st.columns(3)
        selected_academic_year = col1.selectbox("Select Academic Year:", academic_year, key='ac')
        selected_year_level = col2.selectbox("Select Year Level", year_levels, key='yl')
        selected_semester = col3.selectbox("Select Semester", semesters, key='semmy')

        cur.execute("SELECT DISTINCT StudentID FROM academicrecords")
        student_ids = [row[0] for row in cur.fetchall()]
        below_25_gpa = 0
        above_25_gpa = 0
        below_25_cgpa = 0
        above_25_cgpa = 0

        for student_id in student_ids:
            gpa = calculate_gpa(student_id, selected_year_level, selected_semester)
            cgpa = calculate_cgpa(student_id, selected_year_level, selected_semester)
            
            if gpa is not None:
                if gpa > 2.5:
                    below_25_gpa += 1
                else:
                    above_25_gpa += 1
            
            if cgpa is not None:
                if cgpa > 2.5:
                    below_25_cgpa += 1
                else:
                    above_25_cgpa += 1

        rates = calculate_rates(selected_academic_year)
        student_total = rates["student_total"]

        col1, col2 = st.columns(2)
        program_title = col1.text_input("Program Title")
        department = col2.text_input("Department")
        
        col1, col2 = st.columns(2)
        college = col1.text_input("College")
        academic_year = col2.text_input("Academic Year")
        
        col1, col2 = st.columns(2)
        reporting_period = col1.text_input("Reporting Period")
        submission_date = col2.text_input("Report Submission Date")

        st.divider()

        st.write("Program Engagement & Activities")
        objectives = st.text_area("Objectives", height = 200, max_chars=1500)
        co_act = st.text_area("Curricular & Co-Curricular Activities", height = 200, max_chars=1500)
        accomplishments = st.text_area("Accomplishments", height = 200, max_chars=1500)

        st.divider()

        st.write("Program Outputs and Deliverables")
        program_outputs = st.text_area("Program Outputs", height = 200, max_chars=1500)
        deliverables = st.text_area("Deliverables", height = 200, max_chars=1500)

        st.divider()

        st.write("Consultation & Advising")
        date_cons = st.text_area("Date of Consultation", height = 200, max_chars=1500)
        nature_advising = st.text_area("Nature of Advising", height = 200, max_chars=1500)
        action_taken = st.text_area("Action Taken", height = 200, max_chars=1500)

        st.divider()

        risk_challenges = st.text_area("Risks & Challenges", height = 200, max_chars=1500)
        collab_linkages = st.text_area("Collaboration & Linkages", height = 200, max_chars=1500)
        problem_encountered =  st.text_area("Problems Encountered", height = 200, max_chars=1500)
        recom = st.text_area("Recommendations", height = 200, max_chars=1500)
        program_plans = st.text_area("Program Plans", height = 200, max_chars=1500)

        st.divider()

        col1, col2 = st.columns(2)
        prog_adv = col1.text_input("Name of Program Adviser:")
        dept_chairperson = col2.text_input("Department Chairperson:")

        def set_text_properties(paragraph, bold=False, size=11, alignment=None, color=RGBColor(0, 0, 0)):
            for run in paragraph.runs:
                run.font.bold = bold
                run.font.size = Pt(size)
                run.font.color.rgb = color
                if alignment:
                    paragraph.alignment = alignment

        def set_column_width(cell, width):
            cell_width = OxmlElement('w:tcW')
            cell_width.set(qn('w:w'), str(width))
            cell_width.set(qn('w:type'), 'dxa')
            cell._element.get_or_add_tcPr().append(cell_width)

        with st.form("adviser_report_form"):
            submitted = st.form_submit_button("Generate Report")

            if submitted:
                counts = calculate_counts(selected_year_level, selected_semester)
                rates = calculate_rates(selected_academic_year)

                doc = docx.Document()
                
                # Add logo and aligned text in the header
                header = doc.sections[0].header
                header_table = header.add_table(rows=1, cols=2, width=5)

                set_column_width(header_table.columns[0].cells[0], 1000)  # Width in twips (1/20 of a point)
                set_column_width(header_table.columns[1].cells[0], 7000)

                # Add logo to the first cell (adjust path to your logo image)
                logo_cell = header_table.cell(0, 0)
                logo_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture('seal-02.png', width=Inches(0.95))

                # Add text to the second cell
                text_cell = header_table.cell(0,1)
                text_paragraph = text_cell.paragraphs[0]
                text_run1 = text_paragraph.add_run('MSU – ILIGAN INSTITUTE OF TECHNOLOGY\n')
                text_run1.bold = True
                text_run1.font.size = Pt(10)
                text_run1.font.color.rgb = RGBColor(0, 0, 0)

                text_run3 = text_paragraph.add_run('OFFICE OF THE VICE CHANCELLOR FOR ACADEMIC AFFAIRS\n')
                text_run3.font.size = Pt(10)
                text_run3.font.color.rgb = RGBColor(0, 0, 0)

                text_run4 = text_paragraph.add_run('OFFICE OF THE DIRECTOR FOR UNDERGRADUATE PROGRAMS\n')
                text_run4.font.size = Pt(10)
                text_run4.font.color.rgb = RGBColor(0, 0, 0)

                text_run5 = text_paragraph.add_run('Iligan City, Philippines')
                text_run5.font.size = Pt(10)
                text_run5.font.color.rgb = RGBColor(0, 0, 0)

                # Set alignment for text in the second cell
                text_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT 

                headingc = doc.add_heading('ACADEMIC PROGRAM ADVISING PROGRESS REPORT', 0)
                set_text_properties(headingc, size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)


                para1 = doc.add_paragraph()
                run1 = para1.add_run('Program Title: ')
                run1.bold = True
                run1.font.size = Pt(12)
                run1.font.color.rgb = RGBColor(0, 0, 0)
                
                run2 = para1.add_run(program_title)
                run2.bold = False
                run2.font.size = Pt(12)
                run2.font.color.rgb = RGBColor(0, 0, 0)

                # Department
                para2 = doc.add_paragraph()
                run3 = para2.add_run('Department: ')
                run3.bold = True
                run3.font.size = Pt(12)
                run3.font.color.rgb = RGBColor(0, 0, 0)

                run4 = para2.add_run(department)
                run4.bold = False
                run4.font.size = Pt(12)
                run4.font.color.rgb = RGBColor(0, 0, 0)

                # College
                para3 = doc.add_paragraph()
                run5 = para3.add_run('College: ')
                run5.bold = True
                run5.font.size = Pt(12)
                run5.font.color.rgb = RGBColor(0, 0, 0)

                run6 = para3.add_run(college)
                run6.bold = False
                run6.font.size = Pt(12)
                run6.font.color.rgb = RGBColor(0, 0, 0)

                # Academic Year
                para4 = doc.add_paragraph()
                run7 = para4.add_run('Academic Year: ')
                run7.bold = True
                run7.font.size = Pt(12)
                run7.font.color.rgb = RGBColor(0, 0, 0)

                run8 = para4.add_run(academic_year)
                run8.bold = False
                run8.font.size = Pt(12)
                run8.font.color.rgb = RGBColor(0, 0, 0)

                # Reporting Period
                para5 = doc.add_paragraph()
                run9 = para5.add_run('Reporting Period: ')
                run9.bold = True
                run9.font.size = Pt(12)
                run9.font.color.rgb = RGBColor(0, 0, 0)

                run10 = para5.add_run(reporting_period)
                run10.bold = False
                run10.font.size = Pt(12)
                run10.font.color.rgb = RGBColor(0, 0, 0)

                # Report Submission Date
                para6 = doc.add_paragraph()
                run11 = para6.add_run('Report Submission Date: ')
                run11.bold = True
                run11.font.size = Pt(12)
                run11.font.color.rgb = RGBColor(0, 0, 0)

                run12 = para6.add_run(submission_date)
                run12.bold = False
                run12.font.size = Pt(12)
                run12.font.color.rgb = RGBColor(0, 0, 0)

                # Section II: Program Academic Performance Profile
                heading7 = doc.add_heading('I. Program Academic Performance Profile', level=2)
                set_text_properties(heading7, size=12, bold=True)

                # Create table for two-column layout
                table = doc.add_table(rows=8, cols=4)

                # Set column widths (first column wider)
                set_column_width(table.columns[0].cells[0], 5000)  # Width in twips (1/20 of a point)
                set_column_width(table.columns[1].cells[0], 1000)
                set_column_width(table.columns[2].cells[0], 5000)
                set_column_width(table.columns[3].cells[0], 1000)

                # Fill in the table cells
                cells = table.rows[0].cells
                cells[0].text = 'Total Program Enrollees:'
                cells[1].text = str(rates["student_total"])
                cells[2].text = 'Number of Students with INC:'
                cells[3].text = str(counts["inc_count"])

                cells = table.rows[1].cells
                cells[0].text = 'Retention Rate:'
                cells[1].text = f'{rates["retention_rate"]:.2f}%'
                cells[2].text = 'Number of Students withdraw from the program:'
                cells[3].text = str(counts["withdrawn_count"])

                cells = table.rows[2].cells
                cells[0].text = 'Completion Rate:'
                cells[1].text = f'{rates["completion_rate"]:.2f}%'
                cells[2].text = 'Number of Students with failing grades:'
                cells[3].text = str(counts["fail_count"])

                cells = table.rows[3].cells
                cells[0].text = 'Promotion Rate:'
                cells[1].text = f'{rates["promotion_rate"]:.2f}%'
                cells[2].text = 'Number of Rizal Excellence Awardees (1.0 – 1.20):'
                cells[3].text = str(counts["rl_count"])

                cells = table.rows[4].cells
                cells[0].text = 'Failure Rate:'
                cells[1].text = f'{rates["failure_rate"]:.2f}%'
                cells[2].text = 'Number of Chancellor’s Excellence Awardees (1.21 – 1.45):'
                cells[3].text = str(counts["cl_count"])

                cells = table.rows[5].cells
                cells[0].text = 'Dropout Rate:'
                cells[1].text = f'{rates["dropout_rate"]:.2f}%'
                cells[2].text = 'Number of Dean’s Excellence Awardees (1.46 – 1.75):'
                cells[3].text = str(counts["dl_count"])

                cells = table.rows[6].cells
                cells[0].text = 'Average GPA of Students:'
                cells[1].text = f'{average_gpa:.3f}'
                cells[2].text = 'Number of Students with GPA below 2.50:'
                cells[3].text = str(below_25_gpa)

                cells = table.rows[7].cells
                cells[0].text = 'Average CGPA of Students:'
                cells[1].text = f'{average_cgpa:.3f}'
                cells[2].text = 'Number of Students with CGPA below 2.50:'
                cells[3].text = str(below_25_cgpa)

                # Section III: Program Engagement & Activities
                heading8 = doc.add_heading('II. Program Engagement & Activities', level=2)
                set_text_properties(heading8, bold=True, size=12)

                objectives_paragraph = doc.add_paragraph()
                objectives_paragraph.add_run('Objectives: ').bold = True
                objectives_paragraph.add_run(f'{objectives}')
                objectives_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                co_act_paragraph = doc.add_paragraph()
                co_act_paragraph.add_run('Curricular & Co-Curricular Activities: ').bold = True
                co_act_paragraph.add_run(f'{co_act}')
                co_act_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                accomplishments_paragraph = doc.add_paragraph()
                accomplishments_paragraph.add_run('Accomplishments: ').bold = True
                accomplishments_paragraph.add_run(f'{accomplishments}')
                accomplishments_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Section III: Program Engagement & Activities
                heading20 = doc.add_heading('III: Program Outputs and Deliverables', level=2)
                set_text_properties(heading20, bold=True, size=12)

                program_outputs_paragraph = doc.add_paragraph()
                program_outputs_paragraph.add_run('Program Outputs: ').bold = True
                program_outputs_paragraph.add_run(f'{program_outputs}')
                program_outputs_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                deliverables_paragraph = doc.add_paragraph()
                deliverables_paragraph.add_run('Deliverables: ').bold = True
                deliverables_paragraph.add_run(f'{deliverables}')
                deliverables_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Section IV: Consultation & Advising
                heading9 = doc.add_heading('IV. Consultation & Advising', level=2)
                set_text_properties(heading9, bold=True, size=12)

                date_cons_paragraph = doc.add_paragraph()
                date_cons_paragraph.add_run('Date of Consultation: ').bold = True
                date_cons_paragraph.add_run(f'{date_cons}')
                date_cons_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                nature_advising_paragraph = doc.add_paragraph()
                nature_advising_paragraph.add_run('Nature of Advising: ').bold = True
                nature_advising_paragraph.add_run(f'{nature_advising}')
                nature_advising_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                action_taken_paragraph = doc.add_paragraph()
                action_taken_paragraph.add_run('Action Taken: ').bold = True
                action_taken_paragraph.add_run(f'{action_taken}')
                action_taken_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Section V: Risks & Challenges
                heading10 = doc.add_heading('V. Risks & Challenges', level=2)
                set_text_properties(heading10, bold=True, size=12)
                risk_challenges_paragraph = doc.add_paragraph()
                risk_challenges_paragraph.add_run(f'{risk_challenges}')
                risk_challenges_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Ensure to justify the paragraph content

                # Section VI: Collaboration & Linkages
                heading11 = doc.add_heading('VI. Collaboration & Linkages', level=2)
                set_text_properties(heading11, bold=True, size=12)
                collab_linkages_paragraph = doc.add_paragraph()
                collab_linkages_paragraph.add_run(f'{collab_linkages}')
                collab_linkages_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Ensure to justify the paragraph content

                # Section VII: Problems Encountered
                heading12 = doc.add_heading('VII. Problems Encountered', level=2)
                set_text_properties(heading12, bold=True, size=12)
                problem_encountered_paragraph = doc.add_paragraph()
                problem_encountered_paragraph.add_run(f'{problem_encountered}')
                problem_encountered_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Ensure to justify the paragraph content

                # Section VIII: Recommendations
                heading13 = doc.add_heading('VIII. Recommendations', level=2)
                set_text_properties(heading13, bold=True, size=12)
                recom_paragraph = doc.add_paragraph()
                recom_paragraph.add_run(f'{recom}')
                recom_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Ensure to justify the paragraph content

                # Section IX: Program Plans
                heading14 = doc.add_heading('IX. Program Plans', level=2)
                set_text_properties(heading14, bold=True, size=12)
                program_plans_paragraph = doc.add_paragraph()
                program_plans_paragraph.add_run(f'{program_plans}')
                program_plans_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Ensure to justify the paragraph content

                para7 = doc.add_paragraph()
                run13 = para7.add_run('Name of Program Adviser: ')
                run13.bold = True
                run13.font.size = Pt(12)
                run13.font.color.rgb = RGBColor(0, 0, 0)

                run14 = para7.add_run(prog_adv)
                run14.bold = False
                run14.font.size = Pt(12)
                run14.font.color.rgb = RGBColor(0, 0, 0)

                heading16 = doc.add_heading('Signature and Date:', level=2)
                set_text_properties(heading16, bold=True, size=12)

                para8 = doc.add_paragraph()
                run15 = para8.add_run('Department Chairperson:')
                run15.bold = True
                run15.font.size = Pt(12)
                run15.font.color.rgb = RGBColor(0, 0, 0)

                run16 = para8.add_run(dept_chairperson)
                run16.bold = False
                run16.font.size = Pt(12)
                run16.font.color.rgb = RGBColor(0, 0, 0)

                heading18 = doc.add_heading('Signature and Date:', level=2)
                set_text_properties(heading18, bold=True, size=12)

                 # Save the document to a BytesIO object
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

        # Place the download button outside the form submission block
        if submitted:
            st.download_button(
                label="Download Report",
                data=doc_io,
                file_name=f"adviser_report_{selected_year_level}_{selected_semester}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
                